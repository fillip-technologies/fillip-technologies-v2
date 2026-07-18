# -*- coding: utf-8 -*-
"""Generate a detailed DOCX of all website pricing, grouped by category.
Data source: src/data/pricing.ts (the quote-calculator source of truth).
Preserves the site's package -> feature-group -> items structure.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PRIMARY = RGBColor(0x02, 0x42, 0xA2)
DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x64, 0x74, 0x8B)
ACCENT = RGBColor(0x0E, 0x74, 0x90)

GST_RATE = 0.18

def inr(n):
    return "₹" + format(n, ",d")

# --- Pricing data mirrored 1:1 from src/data/pricing.ts (incl. feature groups) ---
CATEGORIES = [
    {
        "name": "Website Development",
        "description": "High-performance, scalable web solutions designed for conversion. "
                       "Priced ~10% below leading market rates.",
        "packages": [
            {"name": "Startup Website", "tagline": "The Starter", "price": 18000, "billing": "one-time",
             "timeline": "3 days", "bestFor": "Small Businesses & Startups (up to 5 pages).",
             "groups": [
                 ("Design & Build", ["Up to 5 Pages — Dynamic Website",
                                      "Logo Design & Slider Functionality",
                                      "Mobile Responsive Design",
                                      "Testimonials & Gallery Sections"]),
                 ("Included Free (1st Year)", ["Domain (.com or .in)",
                                                "Hosting",
                                                "1 Email Account (500 MB)",
                                                "Free SSL Certificate"]),
                 ("Features", ["1 Contact Form",
                                "Live Chat Feature",
                                "Social Profile Links",
                                "Click-to-Call / Email Functionality"]),
             ]},
            {"name": "Small Business Website", "tagline": "The Professional", "price": 31500, "billing": "one-time",
             "timeline": "7 days", "bestFor": "Mid-size Businesses (10 pages).",
             "groups": [
                 ("Design & Build", ["10 Pages — Dynamic Website",
                                      "Logo Design & Slider Functionality",
                                      "Mobile Responsive Design",
                                      "Testimonials & Gallery Sections"]),
                 ("Included Free (1st Year)", ["Domain (.com or .in)",
                                                "Hosting",
                                                "2 Email Accounts (500 MB each)",
                                                "Free SSL Certificate"]),
                 ("Features & Support", ["1 Contact Form & Live Chat",
                                          "Social Profile Links",
                                          "Click-to-Call / Email Functionality",
                                          "Free 4 Hours Support Included"]),
             ]},
            {"name": "Enterprise Website", "tagline": "The E-commerce", "price": 45000, "billing": "one-time",
             "timeline": "15 days", "bestFor": "E-commerce Businesses (up to 50 products).",
             "groups": [
                 ("E-commerce Build", ["Up to 50 Products",
                                        "Payment Gateway Integration",
                                        "Logo Design",
                                        "WordPress Login Details Provided"]),
                 ("Communication", ["3 Contact Forms",
                                     "WhatsApp Chat Integration",
                                     "Social Media Profiles Integration",
                                     "Social Profile Links"]),
                 ("Support", ["Call / Email / Chat Support",
                               "Extra Support Available (hourly add-on)"]),
             ]},
        ],
        "addOns": [
            {"name": "Extra Page", "price": 1000, "billing": "one-time", "unit": "page"},
            {"name": "Content Writing", "price": 1, "billing": "one-time", "unit": "word"},
            {"name": "Hourly Support (beyond included)", "price": 450, "billing": "one-time", "unit": "hour"},
            {"name": "Multilingual Support", "price": 8000, "billing": "one-time"},
            {"name": "Payment Gateway Integration", "price": 6000, "billing": "one-time"},
            {"name": "Hosting & Domain Renewal (from)", "price": 2000, "billing": "one-time"},
        ],
    },
    {
        "name": "Search Engine Optimization (SEO)",
        "description": "Climb the rankings with data-backed, white-hat organic strategies. "
                       "Billed per month. Priced ~10% below leading market rates.",
        "packages": [
            {"name": "Basic SEO", "price": 22500, "billing": "monthly", "timeline": "Per month",
             "bestFor": "Small Business — Max 10 Keywords.",
             "groups": [
                 ("On-Page & Technical", ["Website & Competitor Analysis",
                                           "Content Duplicity Check & Initial Ranking Report",
                                           "Keywords Research",
                                           "Meta Tags, Canonicalization & URL Structure",
                                           "Content & Image Optimization",
                                           "Heading Tag & Website Speed Optimization",
                                           "Robots.txt & Sitemap Creation",
                                           "Google Analytics & Search Console Setup"]),
                 ("Local & Content", ["Blog Optimization (1 post)",
                                       "Google Map Integration & My Business Setup",
                                       "Local Citations (5), Local Classifieds (2)",
                                       "Onsite Blog (1), Offsite Blog (2)"]),
                 ("Off-Page & Reporting", ["Guest Blogging & Broken Backlink Building",
                                            "Resource Link Building, Alerts & Mentions",
                                            "Social Sharing (30), Slide/Podcast/Q&A (1 each)",
                                            "Monthly Ranking, Analytics & Links Report",
                                            "Email, Chat & Call Support"]),
             ]},
            {"name": "Advanced SEO", "price": 36000, "billing": "monthly", "timeline": "Per month",
             "bestFor": "Mid-size Business — Max 20 Keywords.",
             "groups": [
                 ("Everything in Basic, plus", ["Blog Optimization (5 posts)",
                                                 "Local Citations (10), Local Classifieds (5)",
                                                 "Onsite Blog (2), Offsite Blog (4)",
                                                 "Social Sharing (40)",
                                                 "Slide Submissions (2), Podcast (2), Q&A (3)"]),
             ]},
            {"name": "eCommerce SEO", "price": 49500, "billing": "monthly", "timeline": "Per month",
             "bestFor": "eCommerce — Max 40 Keywords.",
             "groups": [
                 ("Everything in Advanced, plus", ["Local Citations (15), Local Classifieds (10)",
                                                    "Onsite Blog (3), Offsite Blog (6)",
                                                    "Social Sharing (50)",
                                                    "Podcast (3), Q&A (5)",
                                                    "Full product & category page optimization"]),
             ]},
        ],
        "addOns": [
            {"name": "Extra Blog Posts (per 4)", "price": 3000, "billing": "monthly"},
            {"name": "Extra Keywords (per 10)", "price": 2500, "billing": "monthly"},
        ],
    },
    {
        "name": "Social Media Management",
        "description": "Building social trust and brand awareness. Billed per month.",
        "packages": [
            {"name": "Growth Package", "price": 10000, "billing": "monthly", "timeline": "Per month",
             "bestFor": "Facebook & Instagram.",
             "groups": [
                 ("Content & Management", ["15–20 Custom Graphic Posts + 6 Reels",
                                            "Hashtag Research & Trending Topic analysis",
                                            "Page Optimization (Bio, CTA, Cover)",
                                            "Community Engagement (comments/DM)"]),
             ]},
            {"name": "Dominance Package", "price": 16000, "billing": "monthly", "timeline": "Per month",
             "bestFor": "Facebook, Instagram, and LinkedIn/Twitter.",
             "groups": [
                 ("Content & Management", ["Daily Posting (30 Posts/Mo) + 10 High-Quality Reels",
                                            "Full Story Strategy (Daily Stories)",
                                            "LinkedIn Industry-Specific Posting",
                                            "Monthly Competitor Social Audit"]),
             ]},
        ],
        "addOns": [
            {"name": "Extra Reels (per 4)", "price": 4000, "billing": "monthly"},
            {"name": "Influencer Outreach", "price": 6000, "billing": "monthly"},
        ],
    },
    {
        "name": "Performance Marketing (Paid Ads)",
        "description": "Targeted campaigns for instant lead generation. Billed per month (ad spend excluded).",
        "packages": [
            {"name": "Professional Ads Management", "price": 16000, "billing": "monthly", "timeline": "Per month",
             "bestFor": "Google Ads (Search/Display) + Meta Ads.",
             "groups": [
                 ("Services", ["Pixel & Conversion API Installation",
                                "A/B Testing for Ad Creatives",
                                "Remarketing Campaigns (past visitors)",
                                "Monthly ROI & Conversion Analysis"]),
             ]},
        ],
        "addOns": [],
    },
    {
        "name": "Software & App Solutions",
        "description": "Building the backbone of your digital infrastructure. Timeline: 20–25 days.",
        "packages": [
            {"name": "Custom ERP/CRM", "price": 120000, "billing": "one-time", "timeline": "20–25 days",
             "bestFor": "Businesses needing automated internal workflows.",
             "groups": [
                 ("Includes", ["Automated workflows",
                                "Employee/Lead management",
                                "Specialized API integrations"]),
             ]},
            {"name": "Mobile Applications", "price": 275000, "billing": "one-time", "timeline": "20–25 days",
             "bestFor": "Native Android and iOS products.",
             "groups": [
                 ("Includes", ["Native Android and iOS development",
                                "Robust backend",
                                "High-load performance and security focus"]),
             ]},
        ],
        "addOns": [
            {"name": "Admin Dashboard", "price": 30000, "billing": "one-time"},
            {"name": "App Maintenance & Support", "price": 8000, "billing": "monthly"},
        ],
    },
]

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

def billing_label(b):
    return "per month" if b == "monthly" else "one-time"

def add_centered(text, size, color, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return p

# --- Cover ---
add_centered("Fillip Technologies Pvt Ltd", 24, PRIMARY, bold=True)
add_centered("Human × Intelligence", 11, GREY, italic=True)
add_centered("Detailed Pricing Guide — Services, Packages & Add-ons", 14, DARK, bold=True)
add_centered("All prices in INR. GST %d%% applicable. Indicative pricing — negotiable on consultation."
             % int(GST_RATE*100), 9, GREY, italic=True)
doc.add_paragraph()

# Table of contents (simple)
toc = doc.add_paragraph()
tr = toc.add_run("Contents")
tr.bold = True
tr.font.size = Pt(12)
tr.font.color.rgb = PRIMARY
for i, cat in enumerate(CATEGORIES, 1):
    doc.add_paragraph("%d.  %s" % (i, cat["name"]), style="List Bullet")
_n = len(CATEGORIES)
doc.add_paragraph("%d.  SMS Communication (Volume-Based Plans)" % (_n + 1), style="List Bullet")
doc.add_paragraph("%d.  SaaS Product Development (Scaling Tiers)" % (_n + 2), style="List Bullet")
doc.add_paragraph("%d.  Custom Packages (Tailored Quote)" % (_n + 3), style="List Bullet")
doc.add_page_break()

for idx, cat in enumerate(CATEGORIES, 1):
    h = doc.add_heading(level=1)
    h.add_run("%d. %s" % (idx, cat["name"])).font.color.rgb = PRIMARY
    dp = doc.add_paragraph()
    rd = dp.add_run(cat["description"])
    rd.italic = True
    rd.font.color.rgb = GREY

    # Comparison table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Package", "Price", "Billing", "Timeline", "Best For"]):
        hdr[i].paragraphs[0].add_run(txt).bold = True
    for pkg in cat["packages"]:
        cells = table.add_row().cells
        pname = pkg["name"] + ((" – " + pkg["tagline"]) if pkg.get("tagline") else "")
        cells[0].paragraphs[0].add_run(pname).bold = True
        pc = cells[1].paragraphs[0].add_run(inr(pkg["price"]))
        pc.bold = True
        pc.font.color.rgb = PRIMARY
        cells[2].paragraphs[0].add_run(billing_label(pkg["billing"]))
        cells[3].paragraphs[0].add_run(pkg.get("timeline", "—"))
        cells[4].paragraphs[0].add_run(pkg.get("bestFor", "—"))
    doc.add_paragraph()

    # Detailed per-package breakdown with feature GROUPS
    for pkg in cat["packages"]:
        ph = doc.add_heading(level=2)
        title = pkg["name"] + ((" — " + pkg["tagline"]) if pkg.get("tagline") else "")
        ph.add_run(title).font.color.rgb = DARK

        # price line
        pl = doc.add_paragraph()
        prun = pl.add_run(inr(pkg["price"]) + " ")
        prun.bold = True
        prun.font.size = Pt(14)
        prun.font.color.rgb = PRIMARY
        brun = pl.add_run("(%s)" % billing_label(pkg["billing"]))
        brun.font.size = Pt(10)
        brun.font.color.rgb = GREY

        meta = doc.add_paragraph()
        m1 = meta.add_run("Timeline: ")
        m1.bold = True
        meta.add_run(pkg.get("timeline", "—") + "     ")
        m2 = meta.add_run("Best for: ")
        m2.bold = True
        meta.add_run(pkg.get("bestFor", "—"))
        for r in meta.runs:
            r.font.size = Pt(9.5)
            if not r.bold:
                r.font.color.rgb = GREY

        for gtitle, items in pkg["groups"]:
            gp = doc.add_paragraph()
            gr = gp.add_run(gtitle)
            gr.bold = True
            gr.font.size = Pt(10.5)
            gr.font.color.rgb = ACCENT
            for it in items:
                doc.add_paragraph(it, style="List Bullet")
        doc.add_paragraph()

    # Add-ons
    if cat["addOns"]:
        ah = doc.add_heading(level=2)
        ah.add_run("Add-ons").font.color.rgb = PRIMARY
        atable = doc.add_table(rows=1, cols=2)
        atable.style = "Light List Accent 1"
        ac = atable.rows[0].cells
        ac[0].paragraphs[0].add_run("Add-on").bold = True
        ac[1].paragraphs[0].add_run("Price").bold = True
        for a in cat["addOns"]:
            row = atable.add_row().cells
            name = a["name"] + ((" — charged per %s" % a["unit"]) if a.get("unit") else "")
            row[0].paragraphs[0].add_run(name)
            suffix = "/month" if a["billing"] == "monthly" else " one-time"
            pr2 = row[1].paragraphs[0].add_run(inr(a["price"]) + suffix)
            pr2.font.color.rgb = PRIMARY
        doc.add_paragraph()

    if idx < len(CATEGORIES):
        doc.add_page_break()

# --- SMS Communication (volume-based plans) ---
doc.add_page_break()
sh = doc.add_heading(level=1)
sh.add_run("%d. SMS Communication (Volume-Based Plans)" % (len(CATEGORIES) + 1)).font.color.rgb = PRIMARY
sd = doc.add_paragraph()
sdr = sd.add_run("Pay-as-you-grow bulk SMS. Pricing scales with monthly volume — the plan tier, "
                 "routing, setup speed and support level are determined by your estimated volume "
                 "(via the live calculator on the site). No setup fee; credits never expire.")
sdr.italic = True
sdr.font.color.rgb = GREY

sms_tiers = [
    ("Starter Plan", "Under 25,000 SMS/mo", "Standard Carrier Route", "Instant Activation", "Standard Email Support"),
    ("Growth Plan", "25,000 – 100,000 SMS/mo", "Standard Direct Route", "Same-Day DLT Setup", "Standard Email & Chat"),
    ("Corporate Plan", "100,000 – 500,000 SMS/mo", "Priority Carrier Route", "Priority DLT Setup", "Priority Phone & Chat"),
    ("Enterprise SLA", "500,000+ SMS/mo", "Dedicated Trunking", "Managed Custom Setup", "24/7 Account Manager"),
]
stable = doc.add_table(rows=1, cols=5)
stable.style = "Light Grid Accent 1"
stable.alignment = WD_TABLE_ALIGNMENT.CENTER
shdr = stable.rows[0].cells
for i, txt in enumerate(["Plan", "Monthly Volume", "Route Type", "Setup Time", "Support Level"]):
    shdr[i].paragraphs[0].add_run(txt).bold = True
for plan, vol, route, setup, support in sms_tiers:
    cells = stable.add_row().cells
    pc = cells[0].paragraphs[0].add_run(plan)
    pc.bold = True
    pc.font.color.rgb = PRIMARY
    cells[1].paragraphs[0].add_run(vol)
    cells[2].paragraphs[0].add_run(route)
    cells[3].paragraphs[0].add_run(setup)
    cells[4].paragraphs[0].add_run(support)
doc.add_paragraph()

inc = doc.add_paragraph()
incr = inc.add_run("All-Inclusive Features (every plan)")
incr.bold = True
incr.font.size = Pt(10.5)
incr.font.color.rgb = ACCENT
for feat in ["DLT Header / Sender ID Registration Support",
             "Unicode SMS Support (Multiple Languages)",
             "Real-Time Live Webhook Delivery Reports",
             "Direct Carrier Route with Smart Failover",
             "No Setup Fee & Credits Never Expire",
             "Standard REST APIs & Developer SDKs"]:
    doc.add_paragraph(feat, style="List Bullet")
smsn = doc.add_paragraph()
smsnr = smsn.add_run("Per-SMS rate is quoted on the site's live volume calculator and confirmed on "
                     "consultation — higher volumes unlock lower per-message rates.")
smsnr.italic = True
smsnr.font.size = Pt(9)
smsnr.font.color.rgb = GREY

# --- SaaS Product Development (scaling tiers) ---
doc.add_page_break()
sah = doc.add_heading(level=1)
sah.add_run("%d. SaaS Product Development (Scaling Tiers)" % (len(CATEGORIES) + 2)).font.color.rgb = PRIMARY
sad = doc.add_paragraph()
sadr = sad.add_run("Custom multi-tenant SaaS engineering. Choose the scaling tier that fits your SaaS "
                   "lifecycle — architecture, capacity, uptime SLA and support scale with your user "
                   "count. Pricing is proposal-based; a standard SaaS MVP takes 8–12 weeks.")
sadr.italic = True
sadr.font.color.rgb = GREY

saas_tiers = [
    ("Startup Tier", "Launch Quickly & Validate", "1-10 Members", "100K API Calls",
     "99% Standard", "Email Support", "Monolith Node Architecture — single server and database, ideal for MVP cost efficiency."),
    ("Scaleup Tier", "Expand Operations & Automate", "10-100 Members", "5M API Calls",
     "99.9% Premium", "Priority Slack", "Decoupled Scaled Cluster — load balancer, Redis cache layer and database read replica."),
    ("Enterprise Tier", "High Security & Compliance", "100+ Members", "100M+ API Calls",
     "99.99% Enterprise", "24/7 Dedicated SLA", "Global High-Availability Mesh — cross-region Kubernetes clusters, WAF shield, globally replicated databases."),
]
for label, subtitle, team, reqs, sla, support, arch in saas_tiers:
    th = doc.add_heading(level=2)
    th.add_run("%s — %s" % (label, subtitle)).font.color.rgb = DARK
    mt = doc.add_table(rows=1, cols=4)
    mt.style = "Light Grid Accent 1"
    mh = mt.rows[0].cells
    for i, txt in enumerate(["Optimal Team Size", "Monthly Requests", "Uptime SLA", "Support"]):
        mh[i].paragraphs[0].add_run(txt).bold = True
    mc = mt.add_row().cells
    mc[0].paragraphs[0].add_run(team)
    mc[1].paragraphs[0].add_run(reqs)
    mc[2].paragraphs[0].add_run(sla)
    mc[3].paragraphs[0].add_run(support)
    ap = doc.add_paragraph()
    apr = ap.add_run("Architecture: ")
    apr.bold = True
    apr.font.color.rgb = ACCENT
    ap.add_run(arch)
    doc.add_paragraph()

# --- Custom Packages / Tailored Quote section ---
doc.add_page_break()
ch = doc.add_heading(level=1)
ch.add_run("%d. Custom Packages (Tailored Quote)" % (len(CATEGORIES) + 3)).font.color.rgb = PRIMARY
cd = doc.add_paragraph()
rcd = cd.add_run("Not sure which package fits? Share your requirement and we craft a tailored quote "
                 "around your goals, budget, and timeline. Choose a project type and an indicative "
                 "budget band below, and our team scopes the exact price.")
rcd.italic = True
rcd.font.color.rgb = GREY

def add_option_block(title, options):
    bp = doc.add_paragraph()
    br = bp.add_run(title)
    br.bold = True
    br.font.size = Pt(10.5)
    br.font.color.rgb = ACCENT
    for o in options:
        doc.add_paragraph(o, style="List Bullet")

add_option_block("Project Types", [
    "New build",
    "Redesign",
    "New feature / module",
    "Maintenance & support",
    "Not sure yet",
])

# Budget bands as a table
bh = doc.add_paragraph()
bhr = bh.add_run("Indicative Budget Bands")
bhr.bold = True
bhr.font.size = Pt(10.5)
bhr.font.color.rgb = ACCENT
btable = doc.add_table(rows=1, cols=1)
btable.style = "Light List Accent 1"
btable.rows[0].cells[0].paragraphs[0].add_run("Budget Range (INR)").bold = True
for band in ["Under ₹25,000", "₹25,000 – ₹50,000", "₹50,000 – ₹1,00,000",
             "₹1,00,000 – ₹3,00,000", "₹3,00,000+"]:
    btable.add_row().cells[0].paragraphs[0].add_run(band).font.color.rgb = PRIMARY

doc.add_paragraph()
add_option_block("Timeline Options", [
    "ASAP",
    "Within 1 month",
    "1–3 months",
    "Flexible",
])

cn = doc.add_paragraph()
cnr = cn.add_run("How it works: pick a project type, budget band and timeline, describe your "
                 "requirement, and our team replies with a tailored quote. No fixed price — scoped "
                 "to your exact needs.")
cnr.italic = True
cnr.font.size = Pt(9)
cnr.font.color.rgb = GREY

# Footer note
doc.add_paragraph()
fn = doc.add_paragraph()
fr = fn.add_run("Note: This is an indicative estimate generated from our standard service catalogue. "
                "Final pricing is negotiable when you visit our office or speak with us over a call. "
                "GST of %d%% applies to all listed prices. Performance-marketing ad spend and third-party "
                "gateway/hosting fees are billed separately where applicable." % int(GST_RATE*100))
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

out = r"D:\Fillip\fillip-technologies\Fillip-Technologies-Pricing.docx"
doc.save(out)
print("Saved:", out)
